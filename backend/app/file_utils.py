import hashlib
import logging
import os
import re
import tempfile

import docx
import fitz
from app.crud import get_document_by_hash
from sqlalchemy.orm import Session
from app.vectorize import extract_chapter_contents_from_bytes, extract_chapter_headings_with_page_numbers_from_bytes

logger = logging.getLogger(__name__)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def hash_content(content):
    if isinstance(content, str):
        content = content.encode("utf-8")
    return hashlib.sha256(content).hexdigest()

def extract_author_from_docx(binary_data: bytes) -> str:
    """Extract author from DOCX core properties, falling back to 'Unknown Author'."""
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(binary_data)
            tmp_path = tmp.name
        doc = docx.Document(tmp_path)
        author = doc.core_properties.author
        return author if author else "Unknown Author"
    except Exception:
        return "Unknown Author"
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

def extract_author_from_pdf(binary_data: bytes) -> str:
    """Extract author from PDF metadata, falling back to 'Unknown Author'."""
    try:
        pdf = fitz.open(stream=binary_data, filetype="pdf")
        metadata = pdf.metadata
        return metadata.get("author") or "Unknown Author"
    except Exception:
        return "Unknown Author"

def extract_title_from_filename(filename: str) -> str:
    """Extract a human-readable title from the uploaded filename."""
    name = os.path.splitext(os.path.basename(filename))[0]
    return name.replace("_", " ").replace("-", " ").strip() or "Unknown Title"

def save_file(doc_hash, file_ext, content):
    filename = f"{doc_hash}.{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    with open(file_path, "wb") as f:
        f.write(content)


async def process_file(file, db: Session):
    content = await file.read()
    file_ext = file.filename.split(".")[-1].lower()

    if file_ext == "pdf":
        doc_author = extract_author_from_pdf(content)
    elif file_ext == "docx":
        doc_author = extract_author_from_docx(content)
    else:
        raise ValueError("Unsupported file type")

    doc_title = extract_title_from_filename(file.filename)
    doc_hash = hash_content(content)

    # Use content hash for reliable duplicate detection
    existing_doc = get_document_by_hash(db, doc_hash)
    if existing_doc:
        logger.info("Document already exists: id=%s", existing_doc.id)
        return {
            "document_id": existing_doc.id,
            "already_exists": True,
            "document_hash": existing_doc.document_hash,
            "title": existing_doc.title,
            "author": existing_doc.author,
            "file_type": existing_doc.file_type,
            "chapters": existing_doc.chapters,
        }

    chapter_headings = extract_chapter_headings_with_page_numbers_from_bytes(content)
    chapters = extract_chapter_contents_from_bytes(content, chapter_headings)
    logger.info("Extracted %d chapters from %s", len(chapters), file.filename)

    for ch in chapters:
        ch["hash"] = hash_content(ch["content"])

    save_file(doc_hash, file_ext, content)

    return {
        "already_exists": False,
        "document_hash": doc_hash,
        "title": doc_title,
        "author": doc_author,
        "file_type": file_ext,
        "chapters": chapters,
    }

def extract_text_from_pdf(binary_data: bytes) -> str:
    """Extract plain text from a PDF file."""
    pdf = fitz.open(stream=binary_data, filetype="pdf")
    return "\n".join(page.get_text() for page in pdf)

def extract_text_from_docx(binary_data: bytes) -> str:
    """Extract plain text from a DOCX file."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(binary_data)
        tmp_path = tmp.name
    try:
        doc = docx.Document(tmp_path)
        return "\n".join(p.text for p in doc.paragraphs)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

def split_into_chapters(text: str):
    # Match section headers like Chapter, Lesson, Unit with numbers or roman numerals
    pattern = re.compile(r"\b(Chapter|Lesson|Unit)\s+([0-9IVXLC]+)\b", re.IGNORECASE)
    matches = list(pattern.finditer(text))

    chapters = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        section_type = match.group(1).capitalize()
        section_number = match.group(2)
        title = f"{section_type} {section_number}"
        content = text[start:end].strip()
        chapters.append({"title": title, "content": content, "section_type": section_type})

    return chapters